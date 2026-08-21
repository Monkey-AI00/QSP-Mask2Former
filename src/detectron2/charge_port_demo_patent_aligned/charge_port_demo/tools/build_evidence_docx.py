from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
WORK_ROOT = PROJECT_ROOT.parent
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK_BLUE = "0B2545"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
RISK_RED = "9B1C1C"
CAUTION = "7A5A00"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


CLAIMS = [
    ("1", "闭环主流程：感知、置信度、主动调整、图谱关系、充电口定位", "src/main.py::run_demo；tests/test_integration.py", "支持（可复现仿真）"),
    ("2", "RGB/深度/点云特征，多尺度注意力融合，分类与关键点热力图双分支", "src/multimodal.py；src/perception.py；tests/test_perception_confidence.py", "支持（结构执行+回放头）"),
    ("3", "车型概率归一化信息熵、热力图峰值平均、加权观测置信度", "src/confidence.py；tests/test_perception_confidence.py", "支持"),
    ("4", "卷积-池化图像分支、MLP-最大池化点云分支、注意力融合与双输出头", "src/multimodal.py::extract_multimodal_features；architecture_trace", "结构级支持"),
    ("5", "类别间隔、热力图回归、KL散度与预测分布方差约束的复合损失", "src/losses.py；tests/test_losses.py", "损失/校准级支持"),
    ("6", "候选位姿、预测置信度、置信度提升量及最大增益选择", "src/active_view.py::select_next_best_view；tests/test_active_view.py", "支持"),
    ("7", "基于位姿变换、部件投影覆盖、视线夹角和遮挡先验预测可观测性", "src/active_view.py::predict_observability；候选评分日志", "支持（几何代理）"),
    ("8", "车型根节点、外观部件观测节点、充电口推理节点及4×4刚体变换边", "src/semantic_graph.py；data/graphs/*_graph.json", "支持"),
    ("9", "像素反投影、手眼变换、部件到充电口刚体变换及多节点融合", "src/geometry.py；src/infer_port.py；tests/test_graph_inference.py", "支持"),
    ("10", "处理器执行方法指令的设备形态", "Python程序可在通用处理器执行；tests/test_integration.py", "程序执行级；实体设备未验证"),
]


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def fmt_xyz(values: list[float] | None) -> str:
    return "—" if values is None else "[" + ", ".join(f"{value:.4f}" for value in values) + "]"


def build_matrix_markdown() -> str:
    lines = [
        "# 权利要求支持矩阵",
        "",
        "> 评估口径：与尚未正式提交的专利文本逐项核对。‘支持’仅指可复现仿真/工程原型的计算链路，不等于真实模型精度或实体机器人验证。",
        "",
        "| 权项 | 核心技术特征 | 对应代码与测试 | 支持状态 | 边界说明 |",
        "|---|---|---|---|---|",
    ]
    boundaries = {
        "1": "三案例闭环均可执行；输入为固定回放与仿真几何。",
        "2": "执行结构与热力图已实现，分类/关键点头由固定回放结果提供先验，不是训练权重。",
        "3": "公式和极值测试已覆盖。",
        "4": "固定权重轻量结构用于证明层级数据流，不构成生产网络架构或精度证明。",
        "5": "损失与确定性校准可执行，尚无真实数据集端到端训练记录。",
        "6": "候选位姿来自元数据/网格接口，未连接真实底盘地图。",
        "7": "用覆盖、角度与遮挡先验作可观测性代理，未执行真实候选视角采集。",
        "8": "运行时按车型装载JSON图谱；尚无在线学习或数据库图谱服务。",
        "9": "矩阵链和多节点融合可执行；手眼矩阵为仿真标定值。",
        "10": "仅验证程序在通用处理器运行，未验证相机、底盘、机械臂和充电枪实体协同。",
    }
    for claim, feature, evidence, status in CLAIMS:
        lines.append(f"| {claim} | {feature} | `{evidence}` | {status} | {boundaries[claim]} |")
    lines += [
        "",
        "## 结论",
        "",
        "权利要求1、3、6—9的主要计算步骤已有可执行代码、单元测试和端到端日志对应；权利要求2、4、5达到结构/损失函数级仿真支持；权利要求10仅有程序执行层支持，实体设备验证仍缺失。",
    ]
    return "\n".join(lines) + "\n"


def build_risk_markdown() -> str:
    return """# 专利提交前技术风险提示

> 本文件是技术一致性诊断，不是法律意见。专利尚未正式提交，建议在代理人定稿时同步核对。

## 高优先级

1. **权利要求4的网络层级表述较硬。** 当前代码执行固定权重的卷积/池化、MLP/最大池化、注意力和双分支头结构，但没有训练后的模型权重或训练日志。若继续保留精确层级限定，建议补充真实训练脚本、数据集版本、权重摘要与训练记录；否则在说明书实施例中明确这是可选结构，避免把仿真结构误写为已训练成品。
2. **权利要求5仍缺真实端到端训练证据。** 已补齐类别间隔损失、热力图回归、KL散度和预测分布方差约束以及固定种子校准，但尚未对真实车辆数据集训练。佐证材料必须使用“损失函数/校准链路已实现”，不得写“模型已完成训练并达到某精度”。
3. **权利要求10的实体设备验证缺失。** 程序可在通用处理器执行，但未联调RGB-D相机、移动底盘、机械臂和充电枪。不得把现有GIF或帧图称为实体机器人充电视频。

## 中优先级

1. 权利要求6中的“局部栅格图”目前由候选位姿配置与网格生成接口模拟，尚未接入SLAM/占据栅格。
2. 权利要求7的候选观测置信度由覆盖度、视线夹角和遮挡先验预测，属于几何可观测性代理，不是移动后真实采集结果。
3. 权利要求8的“动态”体现为按车型运行时选择、校验和可更新关系；当前持久化介质是JSON，不是在线知识图谱数据库。
4. 权利要求9的手眼矩阵和图谱关系为仿真参数；正式实体验证需保存标定板、采集批次、标定残差和矩阵版本。
5. 权利要求10当前引用“权利要求1—8”，未覆盖权利要求9的完整三维变换链，建议请专利代理人核对从属引用范围。

## 已消除的原材料问题

- 删除 `direct_annotation` 真值直通推理；真值只在推理结束后进入 `evaluation.py`。
- 分类置信度由最大概率改为归一化信息熵；定位置信度改为热力图峰值平均。
- 主动视点由远近静态打分改为候选位姿几何可观测性与预测置信度增益排序。
- 二维/三维偏移向量改为经过校验的4×4齐次刚体变换、相机反投影和手眼变换。
- 正式运行不再自动补造缺失输入；派生仿真资产只能通过显式准备接口生成。
- 数据JSON迁移为 `replay_output` 与 `ground_truth` 双区，日志逐步记录来源、坐标系和矩阵。
"""


def build_evidence_markdown(summary: dict[str, Any]) -> str:
    rows = []
    for car_id, case in summary["cases"].items():
        evaluation = case["evaluation"]
        error = "无独立真值" if evaluation["pixel_error"] is None else f'{evaluation["pixel_error"]:.3f} px（人工仿真参考）'
        rows.append(
            f'| {car_id} | {" → ".join(case["view_sequence"])} | {case["final_confidence"]:.3f} | '
            f'{case["prediction_method"]} | {fmt_xyz(case["port_3d_robot"])} | {error} |'
        )
    return f"""# 研发佐证材料（专利对齐仿真版）

## 1. 文档结论

本次诊断对专利权利要求1—10、原研发佐证材料和 `charge_port_demo` 代码进行了逐项对照。修改后的工程已形成可复现的“多模态感知结构—不确定性/置信度—主动视点—动态语义图谱—4×4刚体变换定位—动作输出”闭环。三组案例均从远景开始，经中景、近景后完成定位。

本材料只证明计算方法和工程原型链路可执行。它不证明真实神经网络训练精度、实车泛化能力或实体机器人充电成功率。

## 2. 原材料诊断与修正

原代码存在真值直接返回、候选视角读取未来标注、最大概率代替信息熵、二维偏移代替刚体变换、正式运行自动补造输入等问题。现已完成以下修正：

- `ground_truth` 与 `replay_output` 文件级隔离，真值只进入独立评测模块；
- RGB/深度联合卷积-池化、点云MLP-最大池化、多尺度自注意力和双分支头结构执行；
- 信息熵分类置信度、热力图峰值定位置信度及加权观测置信度；
- 类别间隔、热力图回归、KL散度与预测分布方差约束复合损失；
- 候选位姿覆盖度、视线夹角、遮挡先验、预测置信度和提升量排序；
- 相机内参反投影、手眼4×4变换、部件到充电口4×4关系及加权融合；
- 固定种子、独立案例目录、结构化日志、GIF/MP4后端和SHA-256清单。

## 3. 数据流与防泄漏边界

推理主链：`RGB + Depth → Point Cloud → replay head + multimodal structure → confidence → active view → graph transform fusion → action`。

评测旁路：`final_prediction + ground_truth → evaluation.py`。测试通过替换真值并保持预测完全一致，证明真值不参与感知、置信度、视点规划或充电口定位。

## 4. 权利要求对应

权利要求1、3、6—9的主要计算步骤已有可执行代码、测试和日志；权利要求2、4、5为结构/损失函数级仿真支持；权利要求10仅达到程序可在处理器执行的层级，实体设备尚未验证。详细对应见《权利要求支持矩阵.md》。

## 5. 复现实验

- 固定种子：`{summary["seed"]}`
- 运行命令：`{summary["command"]}`
- Python：`{summary["environment"]["python"]}`
- NumPy：`{summary["environment"]["numpy"]}`
- Pillow：`{summary["environment"]["pillow"]}`
- 置信度阈值：`0.60`

| 案例 | 视点序列 | 最终观测置信度 | 定位方法 | 机器人基座坐标/m | 独立评测 |
|---|---|---:|---|---|---|
{chr(10).join(rows)}

其中，三案例定位方法均为 `graph_rigid_transform_fusion`。`car_A` 和 `car_B` 未提供独立充电口真值，因此不报告误差；`model3` 的0.042像素仅表示图谱参数与人工仿真参考的一致性，不代表真实检测精度。

## 6. 测试与证据文件

- 单元及集成测试：`python -m unittest discover -s tests -v`
- 三案例摘要：`outputs/validation_summary.json`
- 哈希清单：`outputs/manifest_sha256.json`
- 案例日志：`outputs/cases/<car_id>/logs/`
- 逐步帧图：`outputs/cases/<car_id>/frames/`
- 动画：`outputs/cases/<car_id>/demo_video.gif`（安装OpenCV时可输出MP4）

## 7. 阶段性结论与边界

在固定输入和依赖版本下，工程可重复产生相同的结构化定位结果，并能以日志证明各阶段数据来源、坐标系、候选评分和使用矩阵。当前材料足以作为专利计算链路的研发过程佐证，但正式提交前仍应保留“仿真/工程原型”定性，并按《专利提交前技术风险提示.md》补充训练和实体设备证据。
"""


def set_east_asia_font(run, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError("table widths must sum to 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def style_table(table, header: bool = True) -> None:
    table.style = "Table Grid"
    if header:
        for cell in table.rows[0].cells:
            set_cell_shading(cell, LIGHT_GRAY)
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_east_asia_font(run)
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.bold = True


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if "Caption" in styles:
        caption = styles["Caption"]
        caption.font.name = "Calibri"
        caption.font.size = Pt(9)
        caption.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(8)
    if "Evidence Note" not in styles:
        note = styles.add_style("Evidence Note", WD_STYLE_TYPE.PARAGRAPH)
        note.base_style = normal
    note = styles["Evidence Note"]
    note.font.name = "Calibri"
    note.font.size = Pt(9.5)
    note.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    note._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(4)


def configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        header = section.header
        p = header.paragraphs[0]
        p.text = "研发佐证材料｜车辆充电口定位可复现仿真原型"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            set_east_asia_font(run)
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        prefix = fp.add_run("第 ")
        set_east_asia_font(prefix)
        prefix.font.size = Pt(8.5)
        add_field(fp, "PAGE")
        suffix = fp.add_run(" 页 / 共 ")
        set_east_asia_font(suffix)
        suffix.font.size = Pt(8.5)
        add_field(fp, "NUMPAGES")
        suffix2 = fp.add_run(" 页")
        set_east_asia_font(suffix2)
        suffix2.font.size = Pt(8.5)


def add_title(doc: Document, summary: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(68)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("研发佐证材料")
    set_east_asia_font(run)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK_BLUE)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(24)
    run = sub.add_run("一种具身机器人对车辆充电的定位方法及设备｜专利对齐仿真版")
    set_east_asia_font(run)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    table = doc.add_table(rows=4, cols=2)
    values = [
        ("文档性质", "专利提交前研发过程佐证（技术材料，非法律意见）"),
        ("验证范围", "可复现仿真 / 工程原型；不含实体机器人与实车精度验证"),
        ("固定种子", str(summary["seed"])),
        ("验证环境", f'Python {summary["environment"]["python"]}；NumPy {summary["environment"]["numpy"]}；Pillow {summary["environment"]["pillow"]}'),
    ]
    for row, values_row in zip(table.rows, values):
        row.cells[0].text, row.cells[1].text = values_row
        set_cell_shading(row.cells[0], LIGHT_GRAY)
    set_table_geometry(table, [2700, 6660])
    style_table(table, header=False)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("结论摘要：修改后的代码与权利要求1—9的主要计算链条建立了可执行对应；权利要求2、4、5属于结构/损失函数级仿真支持；权利要求10尚缺实体设备验证。")
    set_east_asia_font(run)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    doc.add_page_break()


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None, style: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        first.bold = True
        set_east_asia_font(first)
        rest = p.add_run(text[len(bold_prefix):])
        set_east_asia_font(rest)
    else:
        run = p.add_run(text)
        set_east_asia_font(run)
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_east_asia_font(run)


def add_callout(doc: Document, text: str, color: str = DARK_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = text
    set_cell_shading(table.cell(0, 0), CALLOUT)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    style_table(table, header=False)
    for run in table.cell(0, 0).paragraphs[0].runs:
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = True


def add_claim_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    headers = ["权项", "主要技术特征", "代码 / 测试证据", "结论"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for claim, feature, evidence, status in CLAIMS:
        cells = table.add_row().cells
        for cell, value in zip(cells, [claim, feature, evidence, status]):
            cell.text = value
    set_table_geometry(table, [720, 3300, 3300, 2040])
    style_table(table)


def add_results_table(doc: Document, summary: dict[str, Any]) -> None:
    table = doc.add_table(rows=1, cols=6)
    headers = ["案例", "状态", "视点序列", "最终置信度", "充电口坐标（基座系/m）", "独立评测"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for car_id, case in summary["cases"].items():
        error = case["evaluation"]["pixel_error"]
        eval_text = "无独立真值" if error is None else f"{error:.3f} px*"
        values = [
            car_id,
            case["status"],
            " → ".join(case["view_sequence"]),
            f'{case["final_confidence"]:.3f}',
            fmt_xyz(case["port_3d_robot"]),
            eval_text,
        ]
        for cell, value in zip(table.add_row().cells, values):
            cell.text = value
    set_table_geometry(table, [840, 1140, 1740, 1260, 2700, 1680])
    style_table(table)
    note = doc.add_paragraph(style="Evidence Note")
    note.add_run("* model3 的二维误差仅是图谱参数与人工仿真参考的一致性检查，不代表真实检测精度。car_A、car_B 未提供独立真值，未计算误差。")


def add_figure(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        add_paragraph(doc, f"图像缺失：{path}", style="Evidence Note")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    cp = doc.add_paragraph(caption, style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx(summary: dict[str, Any], manifest: dict[str, Any]) -> Document:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_title(doc, summary)

    doc.add_heading("1. 诊断结论与适用范围", level=1)
    add_paragraph(doc, "本次工作以尚未提交的专利文本、原研发佐证材料及代码雏形为基线，按权利要求1—10逐项核对。修改后的工程已形成从多模态输入、观测置信度、主动调整观测位姿到语义图谱刚体变换定位的闭环，并以测试、日志、帧图和哈希清单留痕。")
    add_callout(doc, "重要边界：本材料证明的是可复现仿真与工程原型，不是训练完成的生产模型、实车精度报告或实体机器人充电验收报告。", CAUTION)
    add_paragraph(doc, "总体判断：", bold_prefix="总体判断：")
    add_bullet(doc, "权利要求1、3、6—9的主要计算步骤已有直接可执行支持。")
    add_bullet(doc, "权利要求2、4、5已有网络结构、热力图、复合损失和校准链路的结构级支持，但没有真实数据集训练权重与训练精度。")
    add_bullet(doc, "权利要求10可证明程序能够在通用处理器执行；相机、底盘、机械臂和充电枪实体协同尚未验证。")

    doc.add_heading("2. 原材料问题与整改结果", level=1)
    table = doc.add_table(rows=1, cols=3)
    for cell, value in zip(table.rows[0].cells, ["原问题", "技术风险", "已完成整改"]):
        cell.text = value
    issues = [
        ("true_port_2d 直接返回", "真值泄漏，无法证明推理", "真值移至 ground_truth，仅 evaluation.py 可读"),
        ("分类置信度取最大概率", "与权利要求3信息熵不符", "采用归一化信息熵置信度并覆盖极值测试"),
        ("未来候选视角标注参与打分", "存在前视信息泄漏", "改为候选位姿几何可观测性和预测增益"),
        ("静态 close > mid > far 规则", "不支持权利要求6—7", "计算覆盖度、视线夹角、遮挡先验和置信度增益"),
        ("二维/三维偏移向量", "不支持刚体变换与手眼链", "改为校验后的4×4齐次矩阵、反投影和融合"),
        ("运行时自动生成缺失输入", "证据来源不可审计", "正式加载只校验输入；派生资产需显式准备"),
    ]
    for issue in issues:
        for cell, value in zip(table.add_row().cells, issue):
            cell.text = value
    set_table_geometry(table, [2760, 2760, 3840])
    style_table(table)

    doc.add_heading("3. 技术架构与数据隔离", level=1)
    add_paragraph(doc, "推理主链：RGB与深度输入 → 深度反投影点云 → 多模态结构与固定回放头 → 信息熵/热力图置信度 → 主动视点 → 图谱刚体变换融合 → 作业动作。")
    add_paragraph(doc, "评测旁路：final_prediction 与 ground_truth 只在定位结束后交给 evaluation.py。防泄漏测试替换评测真值后，两次 final_prediction 完全相同而 evaluation 改变。")
    add_callout(doc, "数据JSON采用 replay_output / ground_truth 双区结构；每一步日志记录数据来源、坐标系、相机到机器人矩阵及候选矩阵。")

    doc.add_heading("4. 与专利技术特征的实现对应", level=1)
    doc.add_heading("4.1 多模态感知与双分支输出（权利要求2、4）", level=2)
    add_paragraph(doc, "图像分支对RGB与归一化深度执行 conv1 → pool1 → conv2 → pool2；点云分支执行 mlp1 → max_pool → mlp2；两个图像尺度与点云令牌经 concatenate → self_attention → conv1x1 融合。全局头执行全局平均池化、全连接和Softmax；局部头按关键点回放中心重建上采样、卷积响应与Sigmoid热力图。")
    add_paragraph(doc, "该实现用于证明结构与数据流可执行；分类与关键点先验来自固定回放头，并非训练权重。运行日志中的 architecture_trace 和 layer_shapes 保存实际执行轨迹。", style="Evidence Note")
    doc.add_heading("4.2 置信度与复合损失（权利要求3、5）", level=2)
    add_paragraph(doc, "分类置信度为1减去归一化信息熵；定位置信度为各关键点热力图峰值平均；观测置信度按 α=0.4、β=0.6 加权。损失模块包括类别间隔Softmax、热力图均方误差、目标分布到预测分布的KL散度及分布方差约束，并提供固定种子温度校准。")
    doc.add_heading("4.3 主动视点（权利要求6、7）", level=2)
    add_paragraph(doc, "候选位姿依据相机到机器人变换将已识别部件投影到候选相机视场，计算视场覆盖、部件法向与视线夹角、遮挡先验，再与分类稳定性融合为预测观测置信度。系统按 confidence_gain 降序选择下一观测点；若最大提升量不为正则停止探索。")
    doc.add_heading("4.4 动态语义图谱与三维定位（权利要求8、9）", level=2)
    add_paragraph(doc, "每个车型图谱含根节点、外观部件观测节点和充电口推理节点；边保存4×4部件到充电口刚体变换。外观部件像素结合深度和相机内参反投影到相机坐标，经 camera_to_robot 手眼矩阵转换到机器人基座坐标，再右乘图谱关系得到充电口候选并按热力图峰值融合。")

    doc.add_heading("5. 权利要求支持矩阵", level=1)
    add_claim_table(doc)
    add_paragraph(doc, "“结构级支持”表示代码具有对应模块、接口、执行轨迹和测试，但不能替代真实训练或设备实验。详细边界另见《权利要求支持矩阵.md》。", style="Evidence Note")

    doc.add_heading("6. 可复现实验设置", level=1)
    settings = doc.add_table(rows=1, cols=2)
    settings.rows[0].cells[0].text = "项目"
    settings.rows[0].cells[1].text = "取值"
    values = [
        ("原型范围", summary["scope"]),
        ("固定种子", str(summary["seed"])),
        ("运行命令", summary["command"]),
        ("Python", summary["environment"]["python"]),
        ("NumPy / Pillow", f'{summary["environment"]["numpy"]} / {summary["environment"]["pillow"]}'),
        ("置信度阈值", "0.60"),
        ("测试命令", "python -m unittest discover -s tests -v"),
        ("哈希文件数", str(manifest.get("file_count", "—"))),
    ]
    for name, value in values:
        row = settings.add_row().cells
        row[0].text, row[1].text = name, value
    set_table_geometry(settings, [2700, 6660])
    style_table(settings)

    doc.add_heading("7. 三案例验证结果", level=1)
    add_results_table(doc, summary)
    add_paragraph(doc, "三组案例均从 view_far 开始，因置信度未达0.60而触发主动探索，依次进入 view_mid 和 view_close，最终状态均为 localized；定位方法全部为 graph_rigid_transform_fusion。")

    doc.add_heading("8. 可视化证据", level=1)
    for index, car_id in enumerate(("car_A", "car_B", "model3"), start=1):
        frame = PROJECT_ROOT / "outputs" / "cases" / car_id / "frames" / "003_view_close_work.png"
        case = summary["cases"][car_id]
        add_figure(doc, frame, f"图8-{index}  {car_id} 近景定位与作业动作（最终观测置信度 {case['final_confidence']:.3f}）")

    doc.add_heading("9. 测试、日志与完整性证据", level=1)
    add_paragraph(doc, "测试覆盖几何往返、非法齐次矩阵、信息熵极值、热力图峰值、多模态防真值泄漏、复合损失确定性、主动视点最大增益、无正增益退出、图谱刚体变换、手眼链路、三案例端到端和汇总一致性。")
    paths = [
        "outputs/validation_summary.json",
        "outputs/manifest_sha256.json",
        "outputs/cases/<car_id>/logs/run_<car_id>_view_far.json",
        "outputs/cases/<car_id>/frames/",
        "tests/",
    ]
    for path in paths:
        add_bullet(doc, path)
    add_paragraph(doc, "SHA-256清单排除缓存、临时渲染文件和ZIP本身，用于核对源代码、测试、数据JSON、README、汇总及案例日志。")

    doc.add_heading("10. 限制、待补证据与提交前建议", level=1)
    risks = [
        "权利要求4：当前为固定权重结构执行与回放头，不是训练后的神经网络；如保留精确层级，应补训练脚本、数据集版本、权重哈希和训练日志。",
        "权利要求5：损失与校准链路已实现，但尚无真实数据集端到端训练证据，不得声称已达到真实识别精度。",
        "权利要求6—7：局部栅格和候选观测由仿真位姿/几何代理实现，尚未接入SLAM和真实底盘。",
        "权利要求8—9：图谱关系和手眼矩阵为仿真参数，实体标定需补标定数据、残差、矩阵版本和重复性实验。",
        "权利要求10：仅有程序执行级支持；实体相机、底盘、机械臂、充电枪与安全联锁均未验证。",
        "请专利代理人核对权利要求10引用‘权利要求1—8’是否需要覆盖权利要求9。",
    ]
    for risk in risks:
        add_bullet(doc, risk)
    add_callout(doc, "提交口径建议：统一使用‘可复现仿真/工程原型已验证计算链路’，不要使用‘实车高精度定位已验证’或‘实体机器人自动充电成功’。", RISK_RED)

    doc.add_heading("11. 阶段性结论", level=1)
    add_paragraph(doc, "经整改，研发佐证材料已不再依赖真值直通或未来标注，代码、测试、日志、帧图和支持矩阵之间建立了可追溯对应。就当前“可复现仿真/工程原型”范围而言，材料能够支持专利权利要求1—9的主要计算思想；对权利要求2、4、5和10的训练/实体层面边界已显式披露。")

    doc.add_heading("附录A 关键文件索引", level=1)
    index_rows = [
        ("感知/融合", "src/multimodal.py；src/perception.py"),
        ("置信度/损失", "src/confidence.py；src/losses.py"),
        ("主动视点", "src/active_view.py"),
        ("几何/图谱/定位", "src/geometry.py；src/semantic_graph.py；src/infer_port.py"),
        ("真值隔离", "src/evaluation.py；tests/test_evaluation_no_leakage.py"),
        ("端到端/汇总", "src/main.py；src/run_validation.py；tests/test_integration.py"),
        ("数据与关系", "data/*/meta.json；data/*/view_*_kpts.json；data/graphs/*_graph.json"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "功能"
    table.rows[0].cells[1].text = "相对路径"
    for left, right in index_rows:
        row = table.add_row().cells
        row[0].text, row[1].text = left, right
    set_table_geometry(table, [2700, 6660])
    style_table(table)
    return doc


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--summary", type=Path, default=PROJECT_ROOT / "outputs" / "validation_summary.json")
    parser.add_argument("--manifest", type=Path, default=PROJECT_ROOT / "outputs" / "manifest_sha256.json")
    parser.add_argument("--markdown", type=Path, default=WORK_ROOT / "研发佐证材料_专利对齐仿真版.md")
    parser.add_argument("--output", type=Path, default=WORK_ROOT / "研发佐证材料_专利对齐仿真版.docx")
    args = parser.parse_args()
    summary = load_json(args.summary)
    manifest = load_json(args.manifest)
    args.markdown.parent.mkdir(parents=True, exist_ok=True)
    args.markdown.write_text(build_evidence_markdown(summary), encoding="utf-8")
    (args.markdown.parent / "权利要求支持矩阵.md").write_text(build_matrix_markdown(), encoding="utf-8")
    (args.markdown.parent / "专利提交前技术风险提示.md").write_text(build_risk_markdown(), encoding="utf-8")
    doc = build_docx(summary, manifest)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
